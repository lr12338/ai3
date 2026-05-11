# -*- coding: utf-8 -*-
"""
为 practices 下每个案例生成「刷题模板.ipynb」：工作任务说明 + 参考结构 + 可运行完整代码。
用法：在本目录执行  python _generate_shuati_templates.py
"""
from __future__ import annotations

import json
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent

NB_META = {
    "kernelspec": {
        "display_name": "Python 3",
        "language": "python",
        "name": "python3",
    },
    "language_info": {
        "name": "python",
        "pygments_lexer": "ipython3",
    },
}


def cell_md(text: str) -> dict:
    return {"cell_type": "markdown", "metadata": {}, "source": [text]}


def cell_code(text: str) -> dict:
    return {
        "cell_type": "code",
        "execution_count": None,
        "metadata": {},
        "outputs": [],
        "source": [text],
    }


def parse_folder_name(folder: Path) -> tuple[str, str]:
    name = folder.name
    if "_" in name:
        pid, title = name.split("_", 1)
        return pid, title
    return name, name


def extract_task_text(html_path: Path) -> str:
    if not html_path.is_file():
        return "_（未找到 exam_preview.html，请参考 README）_"
    html = html_path.read_text(encoding="utf-8", errors="replace")
    m = re.search(
        r"<h3>工作任务</h3>\s*<p>(.*?)</p>\s*<h3>技能要求</h3>",
        html,
        re.DOTALL | re.IGNORECASE,
    )
    if not m:
        return "_（未能解析工作任务，请直接打开 exam_preview.html）_"
    block = m.group(1)
    block = re.sub(r"<br\s*/?>", "\n", block, flags=re.I)
    block = re.sub(r"</p>\s*<p>", "\n\n", block)
    block = re.sub(r"<[^>]+>", "", block)
    block = re.sub(r"[ \t]+\n", "\n", block)
    block = re.sub(r"\n{3,}", "\n\n", block).strip()
    return block


def wrap_bookends(folder: Path, core_cells: list[dict]) -> list[dict]:
    pid, title = parse_folder_name(folder)
    task = extract_task_text(folder / "exam_preview.html")
    head = (
        f"# {pid} {title}\n\n"
        f"## 任务说明（摘自考试预览）\n\n"
        f"{task}\n\n"
        "---\n\n"
        f"以下为**刷题参考模板**：含完整可运行代码（编程题）或答题示例与生成 docx 草稿代码（主观题）。"
    )
    tail = (
        "## 学习提示\n\n"
        "- 编程题：请在**本案例文件夹**下运行 Notebook，以便相对路径 `data/...` 正确。\n"
        "- 主观题：示例答案仅供参考，可按考场要求改写；若未安装 `python-docx`，可先 `pip install python-docx`。\n"
        "- 交卷前核对平台要求的截图、HTML、docx 命名与保存路径。"
    )
    return [cell_md(head)] + core_cells + [cell_md(tail)]


def dump_nb(path: Path, cells: list[dict]) -> None:
    nb = {
        "nbformat": 4,
        "nbformat_minor": 5,
        "metadata": NB_META,
        "cells": cells,
    }
    path.write_text(json.dumps(nb, ensure_ascii=False, indent=1), encoding="utf-8")


# --- 各题核心单元（不含首尾说明） -------------------------------------------------

def cells_1_1_1(_f: Path) -> list[dict]:
    return [
        cell_md(
            "## 数据集说明\n\n"
            "- `PatientID`：患者编号\n"
            "- `Age` / `BMI`：年龄、体质指数\n"
            "- `DaysInHospital`：住院天数（>7 判定高风险）\n"
            "- 其他字段为题面原始字段\n"
        ),
        cell_code(
            "import pandas as pd\n"
            "import numpy as np\n\n"
            "data = pd.read_csv('data/patient_data.csv')\n"
            "data.head()"
        ),
        cell_md("## 任务1：统计住院天数超过 7 天的患者数量及占比，并划分风险等级"),
        cell_code(
            "data['RiskLevel'] = np.where(\n"
            "    data['DaysInHospital'] > 7,\n"
            "    '高风险患者',\n"
            "    '低风险患者',\n"
            ")\n\n"
            "risk_counts = data['RiskLevel'].value_counts()\n"
            "total = len(data)\n"
            "high_risk_ratio = risk_counts['高风险患者'] / total\n"
            "low_risk_ratio = risk_counts['低风险患者'] / total\n\n"
            'print("高风险患者数量:", risk_counts["高风险患者"])\n'
            'print("低风险患者数量:", risk_counts["低风险患者"])\n'
            'print("高风险患者占比:", high_risk_ratio)\n'
            'print("低风险患者占比:", low_risk_ratio)'
        ),
        cell_md("## 任务2：按 BMI 区间统计高风险比例与人数"),
        cell_code(
            "bmi_bins = [0, 18.5, 24, 28, np.inf]\n"
            "bmi_labels = ['偏瘦', '正常', '超重', '肥胖']\n"
            "data['BMIRange'] = pd.cut(\n"
            "    data['BMI'], bins=bmi_bins, labels=bmi_labels, right=False\n"
            ")\n\n"
            "bmi_risk_rate = (\n"
            "    data.groupby('BMIRange')['RiskLevel']\n"
            "    .apply(lambda s: (s == '高风险患者').mean())\n"
            ")\n"
            "bmi_patient_count = data.groupby('BMIRange').size()\n\n"
            'print("BMI区间中高风险患者的比例:")\n'
            "print(bmi_risk_rate)\n"
            'print("各BMI区间患者数:")\n'
            "print(bmi_patient_count)"
        ),
        cell_md("## 任务3：按年龄区间统计高风险比例与人数"),
        cell_code(
            "age_bins = [0, 26, 36, 46, 56, 66, np.inf]\n"
            "age_labels = ['≤25岁', '26-35岁', '36-45岁', '46-55岁', '56-65岁', '＞65岁']\n"
            "data['AgeRange'] = pd.cut(\n"
            "    data['Age'], bins=age_bins, labels=age_labels, right=False\n"
            ")\n\n"
            "age_risk_rate = (\n"
            "    data.groupby('AgeRange')['RiskLevel']\n"
            "    .apply(lambda s: (s == '高风险患者').mean())\n"
            ")\n"
            "age_patient_count = data.groupby('AgeRange').size()\n\n"
            'print("年龄区间中高风险患者的比例:")\n'
            "print(age_risk_rate)\n"
            'print("各年龄区间患者数:")\n'
            "print(age_patient_count)"
        ),
    ]


def cells_1_1_2(_f: Path) -> list[dict]:
    return [
        cell_md(
            "## 数据集说明\n\n"
            "- `SensorID`：传感器 ID\n"
            "- `Timestamp`：时间戳\n"
            "- `SensorType`：Temperature / Humidity / SoilMoisture / SoilPH / Light\n"
            "- `Value`：读数\n"
            "- `Location`：Field1–Field4\n"
        ),
        cell_code(
            "import pandas as pd\n"
            "import numpy as np\n\n"
            "data = pd.read_csv('data/sensor_data.csv')\n"
            "data.head()"
        ),
        cell_md("## 任务1：统计每种传感器的数据数量和平均值"),
        cell_code(
            "sensor_stats = data.groupby('SensorType')['Value'].agg(['count', 'mean'])\n"
            'print("传感器数据数量和平均值:")\n'
            "print(sensor_stats)"
        ),
        cell_md("## 任务2：统计每个位置的温度和湿度传感器平均值"),
        cell_code(
            "location_stats = (\n"
            "    data[data['SensorType'].isin(['Temperature', 'Humidity'])]\n"
            "    .groupby(['Location', 'SensorType'])['Value']\n"
            "    .mean()\n"
            "    .unstack()\n"
            ")\n"
            'print("每个位置的温度和湿度数据平均值:")\n'
            "print(location_stats)"
        ),
        cell_md("## 任务3：异常标记与缺失值处理并导出"),
        cell_code(
            "data['is_abnormal'] = np.where(\n"
            "    ((data['SensorType'] == 'Temperature')\n"
            "     & ((data['Value'] < -10) | (data['Value'] > 50)))\n"
            "    | (\n"
            "        (data['SensorType'] == 'Humidity')\n"
            "        & ((data['Value'] < 0) | (data['Value'] > 100))\n"
            "    ),\n"
            "    True,\n"
            "    False,\n"
            ")\n\n"
            'print("异常值数量:", data["is_abnormal"].sum())\n\n'
            "data['Value'] = data['Value'].ffill().bfill()\n\n"
            "cleaned_data = data.drop(columns=['is_abnormal'])\n"
            "cleaned_data.to_csv('cleaned_sensor_data.csv', index=False)\n"
            'print("数据清洗完成，已保存为 cleaned_sensor_data.csv")'
        ),
    ]


def cells_1_1_3(_f: Path) -> list[dict]:
    return [
        cell_md(
            "## 数据集说明（credit_data.csv）\n\n"
            "含 Age、Income、LoanAmount、CreditScore 等字段，用于完整性与合理性审核。\n"
        ),
        cell_code(
            "import pandas as pd\n"
            "import numpy as np\n"
            "import matplotlib.pyplot as plt\n\n"
            "data = pd.read_csv('data/credit_data.csv')"
        ),
        cell_md("## 任务1：缺失值与重复值统计"),
        cell_code(
            "missing_values = data.isnull().sum()\n"
            "duplicate_values = data.duplicated().sum()\n\n"
            'print("缺失值统计:")\n'
            "print(missing_values)\n"
            'print("重复值统计:")\n'
            "print(duplicate_values)"
        ),
        cell_md("## 任务2：合理性审核字段"),
        cell_code(
            "data['is_age_valid'] = data['Age'].between(18, 70)\n"
            "data['is_income_valid'] = data['Income'] > 0\n"
            "data['is_loan_amount_valid'] = data['LoanAmount'] < (data['Income'] * 5)\n"
            "data['is_credit_score_valid'] = data['CreditScore'].between(300, 850)\n\n"
            "validity_checks = data[\n"
            "    ['is_age_valid', 'is_income_valid', 'is_loan_amount_valid',\n"
            "     'is_credit_score_valid']\n"
            "].all(axis=1)\n"
            "data['is_valid'] = validity_checks\n\n"
            'print("数据合理性检查 describe:")\n'
            "print(\n"
            "    data[\n"
            "        [\n"
            "            'is_age_valid',\n"
            "            'is_income_valid',\n"
            "            'is_loan_amount_valid',\n"
            "            'is_credit_score_valid',\n"
            "            'is_valid',\n"
            "        ]\n"
            "    ].describe()\n"
            ")"
        ),
        cell_md("## 任务3：清洗并保存"),
        cell_code(
            "cleaned_data = data[data['is_valid']].drop(\n"
            "    columns=[\n"
            "        'is_age_valid',\n"
            "        'is_income_valid',\n"
            "        'is_loan_amount_valid',\n"
            "        'is_credit_score_valid',\n"
            "        'is_valid',\n"
            "    ]\n"
            ")\n"
            "cleaned_data.to_csv('cleaned_credit_data.csv', index=False)\n"
            'print("数据清洗完成，已保存为 cleaned_credit_data.csv")'
        ),
    ]


def cells_1_1_4(_f: Path) -> list[dict]:
    return [
        cell_md("## 数据集：`data/user_behavior_data.csv`"),
        cell_code(
            "import pandas\n"
            "import numpy as np\n"
            "import matplotlib.pyplot as plt\n\n"
            "data = pandas.read_csv('data/user_behavior_data.csv')\n"
            'print("数据采集完成，已加载到 DataFrame 中")\n'
            "print(data.head())"
        ),
        cell_md("## 任务：清洗、标准化与导出"),
        cell_code(
            "data = data.dropna()\n\n"
            "data['Age'] = data['Age'].astype(int)\n"
            "data['PurchaseAmount'] = data['PurchaseAmount'].astype(float)\n"
            "data['ReviewScore'] = data['ReviewScore'].astype(int)\n\n"
            "data = data[\n"
            "    (data['Age'].between(18, 70))\n"
            "    & (data['PurchaseAmount'] > 0)\n"
            "    & (data['ReviewScore'].between(1, 5))\n"
            "]\n\n"
            "pa_m, pa_s = data['PurchaseAmount'].mean(), data['PurchaseAmount'].std()\n"
            "rs_m, rs_s = data['ReviewScore'].mean(), data['ReviewScore'].std()\n"
            "data['PurchaseAmount'] = (data['PurchaseAmount'] - pa_m) / pa_s\n"
            "data['ReviewScore'] = (data['ReviewScore'] - rs_m) / rs_s\n\n"
            "data.to_csv('cleaned_user_behavior_data.csv', index=False)\n"
            'print("数据清洗完成，已保存为 cleaned_user_behavior_data.csv")'
        ),
        cell_md("## 任务：分组统计"),
        cell_code(
            "purchase_category_counts = data['PurchaseCategory'].value_counts()\n"
            "gender_purchase_amount_mean = data.groupby('Gender')['PurchaseAmount'].mean()\n\n"
            "bins = [18, 26, 36, 46, 56, 66, np.inf]\n"
            "labels = ['18-25', '26-35', '36-45', '46-55', '56-65', '65+']\n"
            "data['AgeGroup'] = pandas.cut(data['Age'], bins=bins, labels=labels, right=False)\n"
            "age_group_counts = data['AgeGroup'].value_counts().sort_index()\n\n"
            'print("每个购买类别的用户数:\\n", purchase_category_counts)\n'
            'print("不同性别的平均购买金额:\\n", gender_purchase_amount_mean)\n'
            'print("不同年龄段的用户数:\\n", age_group_counts)'
        ),
    ]


def cells_1_1_5(_f: Path) -> list[dict]:
    return [
        cell_md("## 数据集：`data/vehicle_traffic_data.csv`"),
        cell_code(
            "import pandas as pd\n"
            "import numpy as np\n"
            "import matplotlib.pyplot as plt\n\n"
            "data = pd.read_csv('data/vehicle_traffic_data.csv')\n"
            'print("数据采集完成")\n'
            "print(data.head())"
        ),
        cell_md("## 任务：清洗与类型转换"),
        cell_code(
            "data = data.dropna()\n"
            "data['Age'] = data['Age'].astype(int)\n"
            "data['Speed'] = data['Speed'].astype(float)\n"
            "data['TravelDistance'] = data['TravelDistance'].astype(float)\n"
            "data['TravelTime'] = data['TravelTime'].astype(float)\n\n"
            "data = data[\n"
            "    data['Age'].between(18, 70)\n"
            "    & data['Speed'].between(0, 200)\n"
            "    & data['TravelDistance'].between(1, 1000)\n"
            "    & data['TravelTime'].between(1, 1440)\n"
            "]\n\n"
            "data.to_csv('cleaned_vehicle_traffic_data.csv', index=False)\n"
            'print("已保存 cleaned_vehicle_traffic_data.csv")'
        ),
        cell_md("## 审核异常 + 统计"),
        cell_code(
            "mask_ok = (\n"
            "    data['Age'].between(18, 70)\n"
            "    & data['Speed'].between(0, 200)\n"
            "    & data['TravelDistance'].between(1, 1000)\n"
            "    & data['TravelTime'].between(1, 1440)\n"
            ")\n"
            "unreasonable_data = data[~mask_ok]\n"
            'print("不合理的数据:\\n", unreasonable_data.head())\n\n'
            "traffic_event_counts = data['TrafficEvent'].value_counts()\n"
            "gender_stats = data.groupby('Gender')[\n"
            "    ['Speed', 'TravelDistance', 'TravelTime']\n"
            "].mean()\n\n"
            "age_bins = [18, 26, 36, 46, 56, 66, np.inf]\n"
            "age_labels = ['18-25', '26-35', '36-45', '46-55', '56-65', '65+']\n"
            "data['AgeGroup'] = pd.cut(\n"
            "    data['Age'], bins=age_bins, labels=age_labels, right=False\n"
            ")\n"
            "age_group_counts = data['AgeGroup'].value_counts().sort_index()\n\n"
            'print("每种交通事件次数:\\n", traffic_event_counts)\n'
            'print("不同性别统计:\\n", gender_stats)\n'
            'print("年龄段驾驶员数:\\n", age_group_counts)'
        ),
    ]


def cells_2_1_1(_f: Path) -> list[dict]:
    return [
        cell_md("## 数据：`data/auto-mpg.csv`"),
        cell_code(
            "import pandas as pd\n"
            "from sklearn.preprocessing import StandardScaler\n"
            "from sklearn.model_selection import train_test_split\n\n"
            "data = pd.read_csv('data/auto-mpg.csv')\n"
            'print("数据集的前五行:")\n'
            "print(data.head())\n"
            "print(data.dtypes)\n\n"
            'print("\\n检查缺失值:")\n'
            "print(data.isnull().sum())\n"
            "data = data.dropna()\n\n"
            "data['horsepower'] = pd.to_numeric(data['horsepower'], errors='coerce')\n"
            "data = data.dropna(subset=['horsepower'])\n\n"
            "print(data['horsepower'].dtype)\n"
            'print("\\n清洗后缺失值:")\n'
            "print(data.isnull().sum())\n\n"
            "numerical_features = ['displacement', 'horsepower', 'weight', 'acceleration']\n"
            "scaler = StandardScaler()\n"
            "data[numerical_features] = scaler.fit_transform(data[numerical_features])\n\n"
            "selected_features = numerical_features + ['cylinders', 'model year', 'origin']\n"
            "X = data[selected_features]\n"
            "y = data['mpg']\n\n"
            "X_train, X_test, y_train, y_test = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=42\n"
            ")\n\n"
            "cleaned_data = X.copy()\n"
            "cleaned_data['mpg'] = y\n"
            "cleaned_data.to_csv('2.1.1_cleaned_data.csv', index=False)\n"
            'print("\\n已保存 2.1.1_cleaned_data.csv")'
        ),
    ]


def cells_2_1_2(_f: Path) -> list[dict]:
    return [
        cell_md(
            "## 说明\n\n"
            "请将考场发放的 **`大学生低碳生活行为的影响因素数据集.xlsx`** 放到本文件夹后再运行。\n"
            "下列代码按题目常规字段名编写；若列名不一致请对照 `print(data.columns)` 微调。\n"
        ),
        cell_code(
            "import pandas as pd\n"
            "from sklearn.preprocessing import StandardScaler\n"
            "from sklearn.model_selection import train_test_split\n\n"
            "xlsx_path = '大学生低碳生活行为的影响因素数据集.xlsx'\n"
            "data = pd.read_excel(xlsx_path)\n"
            "print(data.head())\n\n"
            "initial_row_count = len(data)\n"
            "data = data.dropna()\n"
            "final_row_count = len(data)\n"
            'print(f"处理后数据行数: {final_row_count}, 删除含缺失行数: {initial_row_count - final_row_count}")\n\n'
            "data = data.drop_duplicates()\n\n"
            "num_col = '4.您的月生活费○≦1,000元   ○1,001-2,000元   ○2,001-3,000元   ○≧3,001元'\n"
            "scaler = StandardScaler()\n"
            "data[[num_col]] = scaler.fit_transform(data[[num_col]])\n\n"
            "feat_cols = [\n"
            "    '1.您的性别○男性 ○女性',\n"
            "    '2.您的年级○大一 ○大二 ○大三 ○大四',\n"
            "    '3.您的生源地○农村 ○城镇（乡镇） ○地县级城市 ○省会城市及直辖市',\n"
            "    num_col,\n"
            "    '5.您进行过绿色低碳的相关生活方式吗?',\n"
            "    '6.您觉得“低碳”，与你的生活关系密切吗？',\n"
            "    '7.低碳生活是否会成为未来的主流生活方式？',\n"
            "    '8.您是否认为低碳生活会提高您的生活质量？',\n"
            "]\n"
            "X = pd.get_dummies(data[feat_cols], drop_first=False)\n"
            "y = data['低碳行为积极性']\n\n"
            "X_train, X_test, y_train, y_test = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=42\n"
            ")\n\n"
            "cleaned_data = pd.concat([X, y], axis=1)\n"
            "cleaned_data.to_csv('2.1.2_cleaned_data.csv', index=False)\n"
            'print("已保存 2.1.2_cleaned_data.csv")'
        ),
    ]


def cells_2_1_3(_f: Path) -> list[dict]:
    return [
        cell_md("## 数据：`data/finance数据集.csv`"),
        cell_code(
            "import pandas as pd\n"
            "import matplotlib.pyplot as plt\n"
            "import seaborn as sns\n"
            "from sklearn.preprocessing import MinMaxScaler\n"
            "from sklearn.model_selection import train_test_split\n\n"
            "data = pd.read_csv('data/finance数据集.csv')\n"
            "print(data.head())\n\n"
            "plt.figure(figsize=(12, 8))\n"
            "numeric_cols = data.select_dtypes(include=['float64', 'int64']).columns\n"
            "for i, col in enumerate(numeric_cols, 1):\n"
            "    plt.subplot(3, 4, i)\n"
            "    sns.boxplot(x=data[col])\n"
            "    plt.title(col)\n"
            "plt.tight_layout()\n"
            "plt.show()\n\n"
            "Q1 = data[numeric_cols].quantile(0.25)\n"
            "Q3 = data[numeric_cols].quantile(0.75)\n"
            "IQR = Q3 - Q1\n\n"
            "mask_out = (\n"
            "    (data[numeric_cols] < (Q1 - 1.5 * IQR))\n"
            "    | (data[numeric_cols] > (Q3 + 1.5 * IQR))\n"
            ").any(axis=1)\n"
            "data_cleaned = data[~mask_out]\n\n"
            "dup_mask = data_cleaned.duplicated()\n"
            "num_dup = dup_mask.sum()\n"
            "data_cleaned = data_cleaned[~dup_mask]\n"
            'print(f"删除的重复行数: {num_dup}")\n\n'
            "scaler = MinMaxScaler()\n"
            "data_cleaned[numeric_cols] = scaler.fit_transform(data_cleaned[numeric_cols])\n\n"
            "target_variable = 'SeriousDlqin2yrs'\n"
            "X = data_cleaned.drop(columns=[target_variable])\n"
            "y = data_cleaned[target_variable]\n\n"
            "X_train, X_test, y_train, y_test = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=42\n"
            ")\n"
            'print(f"训练数据形状: {X_train.shape}")\n'
            'print(f"测试数据形状: {X_test.shape}")\n\n'
            "data_cleaned.to_csv('2.1.3_cleaned_data.csv', index=False)\n"
            'print("已保存 2.1.3_cleaned_data.csv")'
        ),
    ]


def cells_2_1_4(_f: Path) -> list[dict]:
    return [
        cell_md("## 数据：`data/medical_data.csv`"),
        cell_code(
            "import pandas as pd\n"
            "from datetime import datetime\n"
            "from sklearn.preprocessing import MinMaxScaler\n"
            "import matplotlib.pyplot as plt\n"
            "import matplotlib.font_manager as fm\n\n"
            "data = pd.read_csv('data/medical_data.csv', encoding='utf-8')\n"
            "print(data.dtypes)\n"
            "print(data.info())\n"
            "print(data.isnull().sum())\n\n"
            "data['就诊日期'] = pd.to_datetime(data['就诊日期'])\n"
            "data['诊断日期'] = pd.to_datetime(data['诊断日期'])\n"
            "data.rename(columns={'病人ID': '患者ID'}, inplace=True)\n"
            "print(data.head())\n\n"
            "data['诊断延迟'] = (data['诊断日期'] - data['就诊日期']).dt.days\n"
            "data['病程'] = (datetime(2024, 9, 1) - data['诊断日期']).dt.days\n\n"
            "data = data[\n"
            "    (data['诊断延迟'] >= 0)\n"
            "    & (data['病程'] > 0)\n"
            "    & (data['年龄'] < 120)\n"
            "]\n"
            "print(data.describe())\n\n"
            "initial_rows = data.shape[0]\n"
            "data.drop_duplicates(inplace=True)\n"
            'print(f"删除的重复行数: {initial_rows - data.shape[0]}")\n\n'
            "scaler = MinMaxScaler()\n"
            "cols_norm = ['年龄', '体重', '身高']\n"
            "data[cols_norm] = scaler.fit_transform(data[cols_norm])\n"
            "print(data.head())\n\n"
            "dist = data.groupby('疾病类型')['治疗结果'].value_counts().unstack(fill_value=0)\n\n"
            "font_path = 'C:/Windows/Fonts/simhei.ttf'\n"
            "my_font = fm.FontProperties(fname=font_path)\n\n"
            "dist.plot(kind='bar', stacked=True)\n"
            "plt.title('不同疾病类型的治疗结果分布', fontproperties=my_font)\n"
            "plt.xlabel('疾病类型', fontproperties=my_font)\n"
            "plt.ylabel('数量', fontproperties=my_font)\n"
            "plt.xticks(rotation=45, fontproperties=my_font)\n"
            "plt.legend(prop=my_font)\n"
            "plt.tight_layout()\n"
            "plt.show()\n\n"
            "sev_code = pd.Categorical(data['疾病严重程度']).codes\n"
            "plt.scatter(data['年龄'], sev_code)\n"
            "plt.title('年龄与疾病严重程度（编码）', fontproperties=my_font)\n"
            "plt.xlabel('年龄', fontproperties=my_font)\n"
            "plt.ylabel('严重程度编码', fontproperties=my_font)\n"
            "plt.show()\n\n"
            "data.to_csv('2.1.4_cleaned_data.csv', index=False)\n"
            'print("已保存 2.1.4_cleaned_data.csv")'
        ),
    ]


def cells_2_1_5(_f: Path) -> list[dict]:
    return [
        cell_md("## 数据：`data/健康咨询客户数据集.csv`"),
        cell_code(
            "import pandas as pd\n"
            "from sklearn.preprocessing import LabelEncoder\n"
            "from sklearn.model_selection import train_test_split\n"
            "import matplotlib.pyplot as plt\n\n"
            "data = pd.read_csv('data/健康咨询客户数据集.csv')\n"
            "print(data.info())\n"
            "print(data.isnull().sum())\n\n"
            "data_cleaned = data.dropna()\n\n"
            "data_cleaned.loc[:, 'Your age'] = pd.to_numeric(\n"
            "    data_cleaned['Your age'], errors='coerce'\n"
            ")\n"
            "data_cleaned = data_cleaned.dropna(subset=['Your age'])\n"
            "data_cleaned = data_cleaned[data_cleaned['Your age'] >= 0]\n"
            "data_cleaned.loc[:, 'Your age'] = data_cleaned['Your age'].astype(int)\n\n"
            "dup_cnt = data_cleaned.duplicated().sum()\n"
            "data_cleaned = data_cleaned.drop_duplicates()\n"
            'print(f"Removed {dup_cnt} duplicate rows")\n\n'
            "fit_col = 'How do you describe your current level of fitness ?'\n"
            "le = LabelEncoder()\n"
            "data_cleaned[fit_col] = le.fit_transform(data_cleaned[fit_col].astype(str))\n\n"
            "data.columns = data.columns.str.strip()\n"
            "print(list(data.columns))\n\n"
            "data_cleaned = data.dropna(subset=['How often do you exercise?'])\n"
            "counts = data_cleaned['How often do you exercise?'].value_counts()\n"
            "plt.figure(figsize=(10, 6))\n"
            "counts.plot.pie(autopct='%1.1f%%', startangle=90)\n"
            "plt.title('Distribution of Exercise Frequency')\n"
            "plt.ylabel('')\n"
            "plt.show()\n\n"
            "data_filled = data.apply(lambda x: x.fillna(x.mode()[0]))\n"
            "train_data, test_data = train_test_split(\n"
            "    data_filled, test_size=0.2, random_state=42\n"
            ")\n"
            "out = '2.1.5_cleaned_split.csv'\n"
            "pd.concat([train_data, test_data]).to_csv(out, index=False)\n"
            'print("已保存", out)'
        ),
    ]


def cells_2_2_1(_f: Path) -> list[dict]:
    return [
        cell_md("## 信用评分：Logistic + SMOTE"),
        cell_code(
            "import pandas as pd\n"
            "from sklearn.model_selection import train_test_split\n"
            "from sklearn.linear_model import LogisticRegression\n"
            "from sklearn.metrics import classification_report, accuracy_score\n"
            "import pickle\n"
            "from imblearn.over_sampling import SMOTE\n\n"
            "data = pd.read_csv('data/finance数据集.csv')\n"
            "print(data.head())\n\n"
            "X = data.drop(['SeriousDlqin2yrs', 'Unnamed: 0'], axis=1, errors='ignore')\n"
            "y = data['SeriousDlqin2yrs']\n\n"
            "X_train, X_test, y_train, y_test = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=42\n"
            ")\n\n"
            "model = LogisticRegression(max_iter=1000)\n"
            "model.fit(X_train, y_train)\n\n"
            "with open('2.2.1_model.pkl', 'wb') as file:\n"
            "    pickle.dump(model, file)\n\n"
            "y_pred = model.predict(X_test)\n"
            "pd.DataFrame(y_pred, columns=['预测结果']).to_csv(\n"
            "    '2.2.1_results.txt', index=False\n"
            ")\n\n"
            "report = classification_report(y_test, y_pred, zero_division=1)\n"
            "with open('2.2.1_report.txt', 'w') as file:\n"
            "    file.write(report)\n\n"
            "accuracy = accuracy_score(y_test, y_pred)\n"
            'print(f"模型准确率: {accuracy:.4f}")\n\n'
            "smote = SMOTE(random_state=42)\n"
            "X_resampled, y_resampled = smote.fit_resample(X_train, y_train)\n\n"
            "model.fit(X_resampled, y_resampled)\n"
            "y_pred_resampled = model.predict(X_test)\n\n"
            "pd.DataFrame(y_pred_resampled, columns=['预测结果']).to_csv(\n"
            "    '2.2.1_results_xg.txt', index=False\n"
            ")\n\n"
            "report_resampled = classification_report(\n"
            "    y_test, y_pred_resampled, zero_division=1\n"
            ")\n"
            "with open('2.2.1_report_xg.txt', 'w') as file:\n"
            "    file.write(report_resampled)\n\n"
            "accuracy_resampled = accuracy_score(y_test, y_pred_resampled)\n"
            'print(f"重采样后准确率: {accuracy_resampled:.4f}")'
        ),
    ]


def cells_2_2_2(_f: Path) -> list[dict]:
    return [
        cell_md("## 燃油效率：线性管道 + 随机森林"),
        cell_code(
            "import pandas as pd\n"
            "import pickle\n"
            "from sklearn.model_selection import train_test_split\n"
            "from sklearn.linear_model import LinearRegression\n"
            "from sklearn.preprocessing import StandardScaler\n"
            "from sklearn.pipeline import Pipeline\n"
            "from sklearn.ensemble import RandomForestRegressor\n\n"
            "df = pd.read_csv('data/auto-mpg.csv')\n"
            "print(df.head())\n\n"
            "df['horsepower'] = pd.to_numeric(df['horsepower'], errors='coerce')\n"
            "df = df.dropna()\n\n"
            "features = [\n"
            "    'cylinders',\n"
            "    'displacement',\n"
            "    'horsepower',\n"
            "    'weight',\n"
            "    'acceleration',\n"
            "    'model year',\n"
            "    'origin',\n"
            "]\n"
            "X = df[features]\n"
            "y = df['mpg']\n\n"
            "X_train, X_test, y_train, y_test = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=42\n"
            ")\n\n"
            "pipeline = Pipeline(\n"
            "    [('scaler', StandardScaler()), ('linreg', LinearRegression())]\n"
            ")\n"
            "pipeline.fit(X_train, y_train)\n\n"
            "with open('2.2.2_model.pkl', 'wb') as model_file:\n"
            "    pickle.dump(pipeline, model_file)\n\n"
            "y_pred = pipeline.predict(X_test)\n"
            "pd.DataFrame(y_pred, columns=['预测结果']).to_csv(\n"
            "    '2.2.2_results.txt', index=False\n"
            ")\n\n"
            "with open('2.2.2_report.txt', 'w') as results_file:\n"
            "    results_file.write(\n"
            "        f'训练集得分: {pipeline.score(X_train, y_train)}\\n'\n"
            "    )\n"
            "    results_file.write(\n"
            "        f'测试集得分: {pipeline.score(X_test, y_test)}\\n'\n"
            "    )\n\n"
            "rf_model = RandomForestRegressor(n_estimators=100, random_state=42)\n"
            "rf_model.fit(X_train, y_train)\n"
            "y_pred_rf = rf_model.predict(X_test)\n\n"
            "pd.DataFrame(y_pred_rf, columns=['预测结果']).to_csv(\n"
            "    '2.2.2_results_rf.txt', index=False\n"
            ")\n\n"
            "with open('2.2.2_report_rf.txt', 'w') as results_rf_file:\n"
            "    results_rf_file.write(\n"
            "        f'训练集得分: {rf_model.score(X_train, y_train)}\\n'\n"
            "    )\n"
            "    results_rf_file.write(\n"
            "        f'测试集得分: {rf_model.score(X_test, y_test)}\\n'\n"
            "    )"
        ),
    ]


def cells_2_2_3(_f: Path) -> list[dict]:
    return [
        cell_md("## 数据：`data/fitness analysis.csv`"),
        cell_code(
            "import pandas as pd\n"
            "from sklearn.model_selection import train_test_split\n"
            "from sklearn.ensemble import RandomForestRegressor\n"
            "from sklearn.metrics import mean_squared_error, r2_score\n"
            "import pickle\n"
            "import xgboost as xgb\n\n"
            "from pathlib import Path\n\n"
            "_csv = Path('data/fitness analysis.csv')\n"
            "if not _csv.exists():\n"
            "    root = Path('.').resolve().parent\n"
            "    alt = next(root.glob('2.2.3_*/data/fitness analysis.csv'), None)\n"
            "    if alt and alt.is_file():\n"
            "        _csv = alt\n"
            "df = pd.read_csv(_csv)\n"
            "print('数据文件:', _csv)\n"
            "print(df.head())\n\n"
            "df = df.apply(lambda col: col.map(lambda x: x.strip() if isinstance(x, str) else x))\n"
            "df.columns = df.columns.str.strip()\n\n"
            "X = df[\n"
            "    [\n"
            "        'Your gender',\n"
            "        'How important is exercise to you ?',\n"
            "        'How healthy do you consider yourself?',\n"
            "    ]\n"
            "]\n"
            "X = pd.get_dummies(X, drop_first=False)\n\n"
            "age_col = 'Your age'\n"
            "if age_col not in df.columns:\n"
            "    age_col = [c for c in df.columns if c.strip().startswith('Your age')][0]\n\n"
            "y = df[age_col].apply(lambda x: int(str(x).split()[0]))\n\n"
            "X_train, X_test, y_train, y_test = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=42\n"
            ")\n\n"
            "rf_model = RandomForestRegressor(n_estimators=100, random_state=42)\n"
            "rf_model.fit(X_train, y_train)\n\n"
            "with open('2.2.3_model.pkl', 'wb') as model_file:\n"
            "    pickle.dump(rf_model, model_file)\n\n"
            "y_pred = rf_model.predict(X_test)\n"
            "pd.DataFrame(y_pred, columns=['预测结果']).to_csv(\n"
            "    '2.2.3_results.txt', index=False\n"
            ")\n\n"
            "train_score = rf_model.score(X_train, y_train)\n"
            "test_score = rf_model.score(X_test, y_test)\n"
            "mse = mean_squared_error(y_test, y_pred)\n"
            "r2 = r2_score(y_test, y_pred)\n"
            "with open('2.2.3_report.txt', 'w') as report_file:\n"
            "    report_file.write(f'训练集得分: {train_score}\\n')\n"
            "    report_file.write(f'测试集得分: {test_score}\\n')\n"
            "    report_file.write(f'均方误差(MSE): {mse}\\n')\n"
            "    report_file.write(f'决定系数(R^2): {r2}\\n')\n\n"
            "xgb_model = xgb.XGBRegressor(n_estimators=100, random_state=42)\n"
            "xgb_model.fit(X_train, y_train)\n"
            "y_pred_xgb = xgb_model.predict(X_test)\n\n"
            "pd.DataFrame(y_pred_xgb, columns=['预测结果']).to_csv(\n"
            "    '2.2.3_results_xgb.txt', index=False\n"
            ")\n\n"
            "with open('2.2.3_report_xgb.txt', 'w') as xgb_report_file:\n"
            "    xgb_report_file.write(\n"
            "        f'XGBoost训练集得分: {xgb_model.score(X_train, y_train)}\\n'\n"
            "    )\n"
            "    xgb_report_file.write(\n"
            "        f'XGBoost测试集得分: {xgb_model.score(X_test, y_test)}\\n'\n"
            "    )\n"
            "    xgb_report_file.write(\n"
            "        f'XGBoost均方误差(MSE): '\n"
            "        f'{mean_squared_error(y_test, y_pred_xgb)}\\n'\n"
            "    )\n"
            "    xgb_report_file.write(\n"
            "        f'XGBoost决定系数(R^2): '\n"
            "        f'{r2_score(y_test, y_pred_xgb)}\\n'\n"
            "    )"
        ),
    ]


def cells_2_2_4(_f: Path) -> list[dict]:
    return [
        cell_md("## 低碳行为：线性回归（题目数据路径按考场素材调整）"),
        cell_code(
            "import pandas as pd\n"
            "from sklearn.model_selection import train_test_split\n"
            "from sklearn.linear_model import LinearRegression\n"
            "from sklearn.metrics import mean_squared_error, r2_score\n"
            "import pickle\n\n"
            "from pathlib import Path\n\n"
            "_csv = Path('data/fitness analysis.csv')\n"
            "if not _csv.exists():\n"
            "    root = Path('.').resolve().parent\n"
            "    alt = next(root.glob('2.2.3_*/data/fitness analysis.csv'), None)\n"
            "    if alt and alt.is_file():\n"
            "        _csv = alt\n"
            "df = pd.read_csv(_csv)\n"
            "print('数据文件:', _csv)\n"
            "print(df.head())\n\n"
            "df.columns = df.columns.str.strip()\n\n"
            "X = df[\n"
            "    [\n"
            "        'Your gender',\n"
            "        'How important is exercise to you ?',\n"
            "        'How healthy do you consider yourself?',\n"
            "    ]\n"
            "]\n"
            "X = pd.get_dummies(X, drop_first=False)\n\n"
            "age_col = [c for c in df.columns if c.startswith('Your age')][0]\n"
            "y = df[age_col].apply(lambda x: int(str(x).split()[0]))\n\n"
            "X_train, X_test, y_train, y_test = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=42\n"
            ")\n\n"
            "model = LinearRegression()\n"
            "model.fit(X_train, y_train)\n\n"
            "with open('2.2.4_model.pkl', 'wb') as f:\n"
            "    pickle.dump(model, f)\n\n"
            "y_pred = model.predict(X_test)\n"
            "pd.DataFrame(y_pred, columns=['预测结果']).to_csv(\n"
            "    '2.2.4_results.txt', index=False\n"
            ")\n\n"
            "mse = mean_squared_error(y_test, y_pred)\n"
            "r2 = r2_score(y_test, y_pred)\n"
            "with open('2.2.4_report.txt', 'w') as f:\n"
            "    f.write(f'训练集得分: {model.score(X_train, y_train)}\\n')\n"
            "    f.write(f'测试集得分: {model.score(X_test, y_test)}\\n')\n"
            "    f.write(f'MSE: {mse}\\n')\n"
            "    f.write(f'R^2: {r2}\\n')"
        ),
    ]


def cells_2_2_5(_f: Path) -> list[dict]:
    return [
        cell_md("## 步数预测：决策树回归"),
        cell_code(
            "import pandas as pd\n"
            "from sklearn.model_selection import train_test_split\n"
            "from sklearn.tree import DecisionTreeRegressor\n"
            "from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score\n"
            "import pickle\n\n"
            "df = pd.read_csv('data/fitness analysis.csv')\n"
            "print(df.head())\n\n"
            "df.columns = df.columns.str.strip()\n\n"
            "X = df[\n"
            "    [\n"
            "        'Your gender',\n"
            "        'How important is exercise to you ?',\n"
            "        'How healthy do you consider yourself?',\n"
            "    ]\n"
            "]\n"
            "X = pd.get_dummies(X, drop_first=False)\n\n"
            "age_col = [c for c in df.columns if c.startswith('Your age')][0]\n"
            "y = df[age_col].apply(lambda x: int(str(x).split()[0]))\n\n"
            "X_train, X_test, y_train, y_test = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=42\n"
            ")\n\n"
            "tree_model = DecisionTreeRegressor(random_state=42)\n"
            "tree_model.fit(X_train, y_train)\n\n"
            "with open('2.2.5_model.pkl', 'wb') as model_file:\n"
            "    pickle.dump(tree_model, model_file)\n\n"
            "y_pred = tree_model.predict(X_test)\n\n"
            "results = pd.DataFrame({'实际值': y_test, '预测值': y_pred})\n"
            "results.to_csv('2.2.5_results.txt', index=False, sep='\\t')\n\n"
            "mse = mean_squared_error(y_test, y_pred)\n"
            "mae = mean_absolute_error(y_test, y_pred)\n"
            "r2 = r2_score(y_test, y_pred)\n"
            "with open('2.2.5_report.txt', 'w') as f:\n"
            "    f.write(f'均方误差: {mse}\\n')\n"
            "    f.write(f'平均绝对误差: {mae}\\n')\n"
            "    f.write(f'决定系数: {r2}\\n')"
        ),
    ]


def cells_3_2_1(_f: Path) -> list[dict]:
    return [
        cell_md("## ONNX 图像分类 Top5"),
        cell_code(
            "import onnxruntime as ort\n"
            "import numpy as np\n"
            "import scipy.special\n"
            "from PIL import Image\n\n\n"
            "def preprocess_image(\n"
            "    image,\n"
            "    resize_size=256,\n"
            "    crop_size=224,\n"
            "    mean=(0.485, 0.456, 0.406),\n"
            "    std=(0.229, 0.224, 0.225),\n"
            "):\n"
            "    image = image.resize((resize_size, resize_size), Image.BILINEAR)\n"
            "    w, h = image.size\n"
            "    left = (w - crop_size) / 2\n"
            "    top = (h - crop_size) / 2\n"
            "    image = image.crop((left, top, left + crop_size, top + crop_size))\n"
            "    image = np.array(image).astype(np.float32)\n"
            "    image = image / 255.0\n"
            "    image = (image - mean) / std\n"
            "    image = np.transpose(image, (2, 0, 1))\n"
            "    image = image.reshape((1,) + image.shape)\n"
            "    return image\n\n\n"
            "session = ort.InferenceSession(\n"
            "    'ResNet50.onnx', providers=['CPUExecutionProvider']\n"
            ")\n\n"
            "with open('labels.txt', encoding='utf-8') as f:\n"
            "    labels = [line.strip() for line in f.readlines()]\n\n"
            "input_name = session.get_inputs()[0].name\n"
            "output_name = session.get_outputs()[0].name\n\n"
            "image = Image.open('demo.jpg').convert('RGB')\n"
            "processed_image = preprocess_image(image)\n"
            "processed_image = processed_image.astype(np.float32)\n\n"
            "output = session.run(\n"
            "    [output_name], {input_name: processed_image}\n"
            ")[0]\n\n"
            "probabilities = scipy.special.softmax(output, axis=-1)[0]\n"
            "top5_idx = np.argsort(probabilities)[-5:][::-1]\n"
            "top5_prob = probabilities[top5_idx]\n\n"
            'print("Top 5 predicted classes:")\n'
            "for i in range(5):\n"
            "    print(\n"
            "        f'{i+1}: {labels[top5_idx[i]]} - Probability: {top5_prob[i]}'\n"
            "    )"
        ),
    ]


def cells_3_2_2(_f: Path) -> list[dict]:
    return [
        cell_md("## MNIST ONNX"),
        cell_code(
            "import onnxruntime\n"
            "import numpy as np\n"
            "from PIL import Image\n\n"
            "ort_session = onnxruntime.InferenceSession(\n"
            "    'mnist-12.onnx', providers=['CPUExecutionProvider']\n"
            ")\n\n"
            "image = Image.open('demo.jpg').convert('L')\n"
            "image = image.resize((28, 28))\n"
            "image_array = np.array(image, dtype=np.float32)\n"
            "image_array = np.expand_dims(image_array, axis=0)\n"
            "image_array = np.expand_dims(image_array, axis=0)\n\n"
            "ort_inputs = {ort_session.get_inputs()[0].name: image_array}\n"
            "ort_outs = ort_session.run(None, ort_inputs)\n\n"
            "predicted_class = int(np.argmax(ort_outs[0]))\n"
            'print(f"Predicted class: {predicted_class}")'
        ),
    ]


def cells_3_2_3(_f: Path) -> list[dict]:
    return [
        cell_md("## FER 表情 ONNX"),
        cell_code(
            "import numpy as np\n"
            "from PIL import Image\n"
            "import onnxruntime as ort\n\n\n"
            "def preprocess(image_path):\n"
            "    input_shape = (1, 1, 64, 64)\n"
            "    img = Image.open(image_path).convert('L')\n"
            "    img = img.resize((64, 64), Image.Resampling.LANCZOS)\n"
            "    img_data = np.array(img, dtype=np.float32)\n"
            "    img_data = np.expand_dims(img_data, axis=0)\n"
            "    img_data = np.expand_dims(img_data, axis=1)\n"
            "    assert img_data.shape == input_shape\n"
            "    return img_data\n\n\n"
            "emotion_table = {\n"
            "    0: 'Angry',\n"
            "    1: 'Disgust',\n"
            "    2: 'Fear',\n"
            "    3: 'Happy',\n"
            "    4: 'Sad',\n"
            "    5: 'Surprise',\n"
            "    6: 'Neutral',\n"
            "}\n\n"
            "ort_session = ort.InferenceSession(\n"
            "    'emotion-ferplus-8.onnx',\n"
            "    providers=['CPUExecutionProvider'],\n"
            ")\n\n"
            "input_data = preprocess('image.jpg')\n\n"
            "ort_inputs = {ort_session.get_inputs()[0].name: input_data}\n"
            "ort_outs = ort_session.run(None, ort_inputs)\n\n"
            "predicted_label = int(np.argmax(ort_outs[0]))\n"
            "predicted_emotion = emotion_table[predicted_label]\n"
            'print(f"Predicted emotion: {predicted_emotion}")'
        ),
    ]


def cells_3_2_4(_f: Path) -> list[dict]:
    return [
        cell_md("## 花朵识别（单类别概率）"),
        cell_code(
            "import onnxruntime as ort\n"
            "import numpy as np\n"
            "import scipy.special\n"
            "from PIL import Image\n\n\n"
            "def preprocess_image(\n"
            "    image,\n"
            "    resize_size=256,\n"
            "    crop_size=224,\n"
            "    mean=(0.485, 0.456, 0.406),\n"
            "    std=(0.229, 0.224, 0.225),\n"
            "):\n"
            "    image = image.resize((resize_size, resize_size), Image.BILINEAR)\n"
            "    w, h = image.size\n"
            "    left = (w - crop_size) / 2\n"
            "    top = (h - crop_size) / 2\n"
            "    image = image.crop((left, top, left + crop_size, top + crop_size))\n"
            "    image = np.array(image).astype(np.float32)\n"
            "    image = image / 255.0\n"
            "    image = (image - mean) / std\n"
            "    image = np.transpose(image, (2, 0, 1))\n"
            "    image = image.reshape((1,) + image.shape)\n"
            "    return image\n\n\n"
            "session = ort.InferenceSession(\n"
            "    'ResNet50.onnx', providers=['CPUExecutionProvider']\n"
            ")\n\n"
            "with open('labels.txt', encoding='utf-8') as f:\n"
            "    labels = [line.strip() for line in f.readlines()]\n\n"
            "input_name = session.get_inputs()[0].name\n"
            "output_name = session.get_outputs()[0].name\n\n"
            "image = Image.open('demo.jpg').convert('RGB')\n"
            "processed_image = preprocess_image(image).astype(np.float32)\n\n"
            "output = session.run(\n"
            "    [output_name], {input_name: processed_image}\n"
            ")[0]\n\n"
            "accuracy = scipy.special.softmax(output, axis=-1)[0]\n"
            "predicted_idx = int(np.argmax(accuracy))\n"
            "prob_percentage = float(accuracy[predicted_idx]) * 100\n"
            "predicted_label = labels[predicted_idx]\n\n"
            'print(\n'
            '    f"Predicted class: {predicted_label}, Accuracy: {prob_percentage:.2f}%"\n'
            ")"
        ),
    ]


def cells_3_2_5(_f: Path) -> list[dict]:
    return [
        cell_md("## 人脸检测 ONNX（考场目录需含 vision 包与模型）"),
        cell_code(
            "import os\n"
            "import time\n"
            "import cv2\n"
            "import numpy as np\n"
            "import vision.utils.box_utils_numpy as box_utils\n"
            "import onnxruntime as ort\n\n\n"
            "def predict(width, height, confidences, boxes, prob_threshold, iou_threshold=0.3, top_k=-1):\n"
            "    boxes = boxes[0]\n"
            "    confidences = confidences[0]\n"
            "    picked_box_probs = []\n"
            "    picked_labels = []\n"
            "    for class_index in range(1, confidences.shape[1]):\n"
            "        probs = confidences[:, class_index]\n"
            "        mask = probs > prob_threshold\n"
            "        probs = probs[mask]\n"
            "        if probs.shape[0] == 0:\n"
            "            continue\n"
            "        subset_boxes = boxes[mask, :]\n"
            "        box_probs = np.concatenate(\n"
            "            [subset_boxes, probs.reshape(-1, 1)], axis=1\n"
            "        )\n"
            "        box_probs = box_utils.hard_nms(\n"
            "            box_probs, iou_threshold=iou_threshold, top_k=top_k\n"
            "        )\n"
            "        picked_box_probs.append(box_probs)\n"
            "        picked_labels.extend([class_index] * box_probs.shape[0])\n"
            "    if not picked_box_probs:\n"
            "        return np.array([]), np.array([]), np.array([])\n"
            "    picked_box_probs = np.concatenate(picked_box_probs)\n"
            "    picked_box_probs[:, 0] *= width\n"
            "    picked_box_probs[:, 1] *= height\n"
            "    picked_box_probs[:, 2] *= width\n"
            "    picked_box_probs[:, 3] *= height\n"
            "    return (\n"
            "        picked_box_probs[:, :4].astype(np.int32),\n"
            "        np.array(picked_labels),\n"
            "        picked_box_probs[:, 4],\n"
            "    )\n\n\n"
            "class_names = [\n"
            "    name.strip() for name in open('voc-model-labels.txt').readlines()\n"
            "]\n\n"
            "ort_session = ort.InferenceSession(\n"
            "    'version-RFB-320.onnx', providers=['CPUExecutionProvider']\n"
            ")\n"
            "input_name = ort_session.get_inputs()[0].name\n\n"
            "result_path = './detect_imgs_results_onnx'\n"
            "threshold = 0.7\n"
            "path = 'imgs'\n"
            "total_boxes = 0\n\n"
            "if not os.path.exists(result_path):\n"
            "    os.makedirs(result_path, exist_ok=True)\n\n"
            "for file_path in os.listdir(path):\n"
            "    img_path = os.path.join(path, file_path)\n"
            "    orig_image = cv2.imread(img_path)\n"
            "    if orig_image is None:\n"
            "        continue\n"
            "    image = cv2.cvtColor(orig_image, cv2.COLOR_BGR2RGB)\n"
            "    image = cv2.resize(image, (320, 240))\n"
            "    image_mean = np.array([127, 127, 127])\n"
            "    image = (image - image_mean) / 128\n"
            "    image = np.transpose(image, [2, 0, 1])\n"
            "    image = np.expand_dims(image, axis=0)\n"
            "    image = image.astype(np.float32)\n\n"
            "    t0 = time.time()\n"
            "    confidences, boxes = ort_session.run(None, {input_name: image})\n"
            '    print("cost time:{}".format(time.time() - t0))\n\n'
            "    boxes, labels, probs = predict(\n"
            "        orig_image.shape[1],\n"
            "        orig_image.shape[0],\n"
            "        confidences,\n"
            "        boxes,\n"
            "        threshold,\n"
            "    )\n"
            "    for i in range(boxes.shape[0]):\n"
            "        box = boxes[i, :]\n"
            "        label = f'{class_names[labels[i]]}: {probs[i]:.2f}'\n"
            "        cv2.rectangle(\n"
            "            orig_image,\n"
            "            (box[0], box[1]),\n"
            "            (box[2], box[3]),\n"
            "            (255, 255, 0),\n"
            "            4,\n"
            "        )\n"
            "        cv2.imwrite(os.path.join(result_path, file_path), orig_image)\n"
            "    total_boxes += boxes.shape[0]\n\n"
            'print("sum:{}".format(total_boxes))'
        ),
    ]


def cells_doc_style(_folder: Path, doc_name: str, sample_md_parts: list[str]) -> list[dict]:
    """主观题 / 大纲题：Markdown 参考范文 + 可选 python-docx 导出。"""
    body = "\n\n".join(sample_md_parts)
    paras_repr = repr(sample_md_parts)
    code = (
        "# 将参考答案段落写入 docx（正式交卷文件名以 exam_preview / 素材为准）\n"
        "from pathlib import Path\n\n"
        "try:\n"
        "    from docx import Document\n"
        "except ImportError:\n"
        "    print('可选安装: pip install python-docx')\n"
        "else:\n"
        f"    paragraphs = {paras_repr}\n"
        "    doc = Document()\n"
        "    for para in paragraphs:\n"
        "        doc.add_paragraph(para)\n"
        f"    out = Path({doc_name!r})\n"
        "    doc.save(out)\n"
        '    print("已生成草稿:", out.resolve())'
    )
    return [
        cell_md("## 参考答案示例（可自行改写）\n\n" + body),
        cell_code(code),
    ]


def subjective(folder: Path, doc_name: str, samples: list[str]) -> list[dict]:
    return cells_doc_style(folder, doc_name, samples)


# --- registry -----------------------------------------------------------------

def cells_3_1_product(
    xlsx_name: str, doc_report: str, doc_opt: str, dim_lines: tuple[str, str, str]
) -> list[dict]:
    """智能产品数据分析类：读 xlsx → 自动摘要 → 生成两道 docx 草稿。"""
    d1, d2, d3 = dim_lines
    md = (
        "## 数据分析与 docx 输出（考场素材）\n\n"
        f"- 请将 **`{xlsx_name}`** 放在本案例文件夹下。\n"
        f"- 报告：`{doc_report}`；优化方案：`{doc_opt}`。\n\n"
        f"分析维度提示：\n1. {d1}\n2. {d2}\n3. {d3}\n"
    )
    code = (
        "import pandas as pd\n"
        "from pathlib import Path\n\n"
        f"xlsx = Path({xlsx_name!r})\n"
        "if not xlsx.exists():\n"
        "    raise FileNotFoundError(\n"
        "        f'请将数据集复制到: {{xlsx.resolve().parent}}'\n"
        "    )\n\n"
        "df = pd.read_excel(xlsx)\n"
        'print("形状:", df.shape)\n'
        "print(df.head())\n"
        "print(df.describe(include='all'))\n\n"
        "obj_cols = df.select_dtypes(include=['object']).columns.tolist()\n"
        "num_cols = df.select_dtypes(include=['number']).columns.tolist()\n"
        "lines = []\n"
        "lines.append('# 分析报告（草稿，请按实际列名改写）')\n"
        "for title, cols in [('分类字段概览', obj_cols), ('数值字段概览', num_cols)]:\n"
        "    lines.append(f'## {title}')\n"
        "    for c in cols[:12]:\n"
        "        lines.append(f'### {c}')\n"
        "        if c in obj_cols:\n"
        "            vc = df[c].astype(str).value_counts().head(8)\n"
        "            lines.append(vc.to_string())\n"
        "        else:\n"
        "            lines.append(df[c].describe().to_string())\n"
        "report_body = '\\n\\n'.join(lines)\n"
        "opt_body = (\n"
        "    '# 优化方向（示例）\\n\\n'\n"
        "    '1. 响应链路：缩短解析与设备回调耗时，加入超时重试与本地缓存。\\n'\n"
        "    '2. 个性化：基于高频场景做快捷入口与默认策略，减少操作步数。\\n'\n"
        "    '3. 稳定性：对异常值与弱网场景做降级策略，完善日志与远程诊断。\\n'\n"
        ")\n\n"
        "try:\n"
        "    from docx import Document\n"
        "except ImportError:\n"
        "    print('请安装: pip install python-docx openpyxl pandas')\n"
        "else:\n"
        "    def write_simple_doc(path: Path, text: str) -> None:\n"
        "        doc = Document()\n"
        "        for block in text.split('\\n\\n'):\n"
        "            block = block.strip()\n"
        "            if not block:\n"
        "                continue\n"
        "            if block.startswith('# '):\n"
        "                doc.add_heading(block[2:].strip(), level=1)\n"
        "            elif block.startswith('## '):\n"
        "                doc.add_heading(block[3:].strip(), level=2)\n"
        "            elif block.startswith('### '):\n"
        "                doc.add_heading(block[4:].strip(), level=3)\n"
        "            else:\n"
        "                doc.add_paragraph(block)\n"
        "        doc.save(path)\n\n"
        f"    write_simple_doc(Path({doc_report!r}), report_body)\n"
        f"    write_simple_doc(Path({doc_opt!r}), opt_body)\n"
        f'    print("已生成:", {doc_report!r}, {doc_opt!r})'
    )
    return [cell_md(md), cell_code(code)]


BUILDERS = {
    "1.1.1": cells_1_1_1,
    "1.1.2": cells_1_1_2,
    "1.1.3": cells_1_1_3,
    "1.1.4": cells_1_1_4,
    "1.1.5": cells_1_1_5,
    "2.1.1": cells_2_1_1,
    "2.1.2": cells_2_1_2,
    "2.1.3": cells_2_1_3,
    "2.1.4": cells_2_1_4,
    "2.1.5": cells_2_1_5,
    "2.2.1": cells_2_2_1,
    "2.2.2": cells_2_2_2,
    "2.2.3": cells_2_2_3,
    "2.2.4": cells_2_2_4,
    "2.2.5": cells_2_2_5,
    "3.1.1": lambda _f: cells_3_1_product(
        "智能音箱数据集.xlsx",
        "3.1.1-1.docx",
        "3.1.1-2.docx",
        (
            "用户使用习惯：哪些功能最常被使用",
            "功能使用频率：热门与冷门功能",
            "响应时间：各功能平均耗时与瓶颈",
        ),
    ),
    "3.1.2": lambda _f: cells_3_1_product(
        "智能照明系统数据集.xlsx",
        "3.1.2-1.docx",
        "3.1.2-2.docx",
        (
            "用户使用习惯：不同时段对亮度与颜色的偏好",
            "功能使用频率：常用与少用场景",
            "响应时间：指令到调光完成的延迟",
        ),
    ),
    "3.1.3": lambda _f: cells_3_1_product(
        "智能健康手环数据集.xlsx",
        "3.1.3-1.docx",
        "3.1.3-2.docx",
        (
            "用户活动模式：一周内不同时段活动水平",
            "健康指标关注度：步数、心率、睡眠等使用热度",
            "数据同步性能：手环与 App 传输延迟",
        ),
    ),
    "3.1.4": lambda _f: cells_3_1_product(
        "智能健康监测系统数据集.xlsx",
        "3.1.4-1.docx",
        "3.1.4-2.docx",
        (
            "用户活动周期：一天中健康指标变化",
            "健康指标偏好：常用与少用监测项",
            "系统响应与准确性：响应时间与误报风险",
        ),
    ),
    "3.1.5": lambda _f: cells_3_1_product(
        "智能家居环境控制系统数据集.xlsx",
        "3.1.5-1.docx",
        "3.1.5-2.docx",
        (
            "用户环境偏好：不同时段温湿度光照偏好",
            "系统响应时间：操作到反馈延迟",
            "能源消耗：平均能耗与节能潜力",
        ),
    ),
    "3.2.1": cells_3_2_1,
    "3.2.2": cells_3_2_2,
    "3.2.3": cells_3_2_3,
    "3.2.4": cells_3_2_4,
    "3.2.5": cells_3_2_5,
}


def default_subjective_cells(folder: Path) -> list[dict]:
    pid = parse_folder_name(folder)[0]
    title = parse_folder_name(folder)[1]
    return subjective(
        folder,
        f"{pid}-answer-draft.docx",
        [
            f"（{pid} {title}）参考答案骨架：",
            "一、现状概述：简要复述业务背景与当前模块目标。",
            "二、主要问题：列出 2～3 条痛点，每条说明对用户体验或业务指标的影响。",
            "三、优化措施：针对每条问题给出可落地的流程/模型/交互改进。",
            "四、预期效果：用可验证表述（准确率、响应时间、满意度等）收尾。",
        ],
    )


def build_cells(folder: Path) -> list[dict]:
    pid = parse_folder_name(folder)[0]
    if pid in BUILDERS:
        return BUILDERS[pid](folder)
    if pid.startswith("1.2."):
        return default_subjective_cells(folder)
    if pid.startswith("4."):
        return subjective(
            folder,
            f"{pid}-outline.docx",
            [
                "培训/指导大纲参考结构：",
                "一、培训目标：对象、学时、完成后能力描述。",
                "二、课程内容：分模块列出知识点与课时。",
                "三、实操环节：工具环境、练习题目、评分要点。",
                "四、考核方式：理论/实操比例与通过标准。",
                "五、附录：参考资料与常见问题。",
            ],
        )
    return default_subjective_cells(folder)


def main() -> None:
    for folder in sorted(ROOT.iterdir()):
        if not folder.is_dir() or folder.name.startswith("_"):
            continue
        out = folder / "刷题模板.ipynb"
        cells = wrap_bookends(folder, build_cells(folder))
        dump_nb(out, cells)
        print("written:", out.relative_to(ROOT))


if __name__ == "__main__":
    main()
